#!/usr/bin/env python3
"""Erzeugt zwei Familienrechts-Testakten zum Unterhalt.

Die Akten sind bewusst als Arbeitsakten gebaut: Markdown-Stücke, EMLs, DOCX,
XLSX, CSV und PDF-Anlagen entstehen aus einer konsistenten Falllogik. Danach
kann scripts/build-testakte-gesamt-pdf.py die Gesamt-PDFs erzeugen.
"""

from __future__ import annotations

import csv
import textwrap
from email.message import EmailMessage
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent.parent
TESTAKTEN = ROOT / "testakten"


def clean(text: str) -> str:
    return textwrap.dedent(text).strip() + "\n"


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(clean(text), encoding="utf-8")


def write_csv(path: Path, rows: list[list[object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerows(rows)


def write_eml(path: Path, subject: str, sender: str, to: str, date: str, body: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = to
    msg["Date"] = date
    msg.set_content(clean(body))
    path.write_bytes(bytes(msg))


def style_doc(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.bold = True


def write_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    doc.add_heading(title, level=1)
    for heading, paras in sections:
        doc.add_heading(heading, level=2)
        for para in paras:
            doc.add_paragraph(para)
    doc.save(path)


def write_pdf(path: Path, title: str, blocks: list[tuple[str, str | list[list[str]]]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("CaseTitle", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#174B4F"))
    h2 = ParagraphStyle("CaseH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.black)
    body = ParagraphStyle("CaseBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13)
    story = [Paragraph(title, h1), Spacer(1, 0.3 * cm)]
    for heading, content in blocks:
        story.append(Paragraph(heading, h2))
        if isinstance(content, list):
            tbl = Table([[Paragraph(str(c), body) for c in row] for row in content], repeatRows=1, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF4F4")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C6C6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
        else:
            for para in clean(content).split("\n\n"):
                story.append(Paragraph(para.replace("\n", " "), body))
        story.append(Spacer(1, 0.22 * cm))
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=1.8 * cm, rightMargin=1.8 * cm, topMargin=1.7 * cm, bottomMargin=1.7 * cm)
    doc.build(story)


def workbook_style(ws) -> None:
    header_fill = PatternFill("solid", fgColor="DDEEEE")
    thin = Side(style="thin", color="C6D2D2")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        cell.border = Border(bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[chr(64 + col)].width = min(42, max(12, max(len(str(ws.cell(row=r, column=col).value or "")) for r in range(1, ws.max_row + 1)) + 2))


def build_stufenklage_case() -> None:
    slug = "unterhalt-stufenklage-selbststaendiger-hamburg"
    base = TESTAKTEN / slug
    for d in ["emails", "docx", "xlsx", "pdfs", "csv", "gesamt-pdf"]:
        (base / d).mkdir(parents=True, exist_ok=True)

    write_text(base / "README.md", """
    # Akte: Unterhalt Lene Harms ./. Joris Harms — Auskunft, Stufenklage, Selbstständigen-Einkommen

    **Arbeitsakte.** Alle Personen, Anschriften, Schulen, Unternehmen und Aktenzeichen sind anonymisiert. Die Akte gehört fachlich zum Plugin `fachanwalt-familienrecht`.

    ## Kurzbild

    Lene Harms lebt seit Januar 2026 mit den Kindern Tilda (14) und Bent (8) getrennt von Joris Harms in Hamburg-Winterhude. Joris führt die Harms Lichtbau GmbH und ein Nebengewerbe für Veranstaltungstechnik. Er zahlt seit März nur noch 900 EUR monatlich und behauptet, sein Einkommen sei wegen Auftragsrückgangs eingebrochen. Gleichzeitig liegen Hinweise auf private Firmenwagennutzung, Gesellschafterdarlehen, Barentnahmen, neue Kameraausstattung und eine nicht erklärte Vermietung einer Ferienwohnung in Grömitz vor.

    Die Akte ist auf Auskunft, vorläufigen Unterhalt und Stufenklage angelegt. Die Einkommensunterlagen sind widersprüchlich und unvollständig; der Rechner enthält nur eine Schätzung. Das System soll erkennen, welche Auskünfte und Belege fehlen, wie der Auskunftsanspruch aufzubauen ist und welche Beträge vorläufig verlangt werden können.

    ## Aktenstücke

    | Nr. | Datei | Inhalt |
    |---|---|---|
    | 01 | `01_mandatsnotiz_trennung_unterhalt.md` | Erstgespräch, Trennung, Kinder, bisherige Zahlungen |
    | 02 | `02_familien_und_kinderdaten.md` | Kinder, Betreuung, Schule, Mehrbedarf |
    | 03 | `03_zahlungsfluss_und_verzug.md` | Zahlungsverlauf, Inverzugsetzung, offene Beträge |
    | 04 | `04_einkommen_lenes_arbeitgeberbescheinigung.md` | Lohn, Steuerklasse, Krankenversicherung, Fahrtkosten |
    | 05 | `05_joris_unvollstaendige_auskunft.md` | Teilunterlagen, Widersprüche, fehlende Jahre |
    | 06 | `06_bwa_fragmente_firmenwagen_entnahmen.md` | BWA-Auszüge, Tesla, Privatanteile, Darlehen |
    | 07 | `07_ferienwohnung_groemitz_hinweise.md` | Vermietungsindizien, Plattform-Screenshots als Notiz |
    | 08 | `08_auskunftsaufforderung_entwurf.md` | Entwurf Auskunfts- und Beleganforderung |
    | 09 | `09_stufenantrag_entwurf.md` | Entwurf Stufenantrag mit Auskunft, Versicherung, Zahlung |
    | 10 | `10_einstweilige_anordnung_vorlaeufiger_unterhalt.md` | EA-Strategie und vorläufige Bezifferung |
    | 11 | `11_belegmatrix_luecken_und_beweiswert.md` | Lückenliste mit Rechtsfolge |
    | 12 | `12_whatsapp_elternchat_unterhalt.md` | Auszug Elternchat zu Geld und Betreuung |

    ## Anhänge

    | Datei | Inhalt |
    |---|---|
    | `emails/2026-02-16_lenes_erstanfrage.eml` | Erstanfrage an Kanzlei |
    | `emails/2026-03-08_joris_kuerzt_zahlung.eml` | Ankündigung der Kürzung |
    | `docx/auskunftsaufforderung_joris_harms.docx` | Aufforderung nach Auskunft und Belegen |
    | `docx/stufenantrag_unterhalt_harms_entwurf.docx` | Entwurf Stufenantrag |
    | `xlsx/vorlaeufige_unterhaltsschaetzung_harms.xlsx` | Schätzung mit Lückenmarkierung |
    | `csv/kontoauszug_unterhalt_zahlungen.csv` | Zahlungen Januar bis Juni 2026 |
    | `pdfs/bwa_auszug_harms_lichtbau_q1_2026.pdf` | Auszug BWA und Entnahmen |
    | `pdfs/arbeitgeberbescheinigung_lene_harms.pdf` | Arbeitgeberbescheinigung Lene |
    """)

    write_text(base / "01_mandatsnotiz_trennung_unterhalt.md", """
    # Mandatsnotiz vom 04.06.2026

    Lene Harms erschien um 09:15 Uhr mit einem Leitzordner, zwei ausgedruckten Kontoauszügen und mehreren E-Mails ihres Ehemannes. Sie lebt seit dem 18.01.2026 getrennt. Joris Harms ist nach einem Streit über die Firmenfinanzen aus der Ehewohnung in der Ohlsdorfer Straße ausgezogen und wohnt nach ihrer Darstellung in einer möblierten Wohnung in Hamburg-Barmbek. Die Kinder Tilda und Bent wohnen bei Lene. Joris holt Bent meistens mittwochs zum Fußballtraining ab und nimmt beide Kinder an jedem zweiten Wochenende, wobei Tilda seit April öfter absagt.

    Lene arbeitet 30 Stunden pro Woche als Projektkoordinatorin bei der Nordkai Museum Service gGmbH. Ihr durchschnittliches Nettoeinkommen soll nach den letzten drei Gehaltsabrechnungen 2.386,40 EUR betragen. Sie zahlt die Warmmiete der Ehewohnung von 1.780,00 EUR, Hortkosten für Bent von 158,00 EUR und Nachhilfe für Tilda von 185,00 EUR. Kindergeld erhält sie auf ihr Konto.

    Joris zahlte im Februar 2026 noch 2.250,00 EUR, im März 1.500,00 EUR und seit April nur noch 900,00 EUR. Er schreibt, er könne wegen Auftragsflaute nichts anderes leisten. Lene hält das für vorgeschoben, weil Joris im April eine neue Kameraausrüstung gekauft habe und weiterhin mit einem Tesla Model Y aus der GmbH fahre. Sie berichtet außerdem, Joris habe die Ferienwohnung seiner Mutter in Grömitz "praktisch übernommen" und vermiete sie über ein Portal, die Einnahmen liefen aber nicht über sein privates Konto.

    Ziel des Mandats ist zunächst nicht die Scheidung, sondern die Sicherung laufenden Kindes- und Trennungsunterhalts. Lene möchte keine Eskalation im Umgang, ist aber bereit, Auskunft gerichtlich geltend zu machen. Sie hat Angst, mit der Miete in Rückstand zu geraten.
    """)

    write_text(base / "02_familien_und_kinderdaten.md", """
    # Familien- und Kinderdaten

    | Person | Geburtsdatum | Rolle | Aktuelle Situation |
    |---|---:|---|---|
    | Lene Harms | 17.03.1984 | Mandantin | 30 Std./Woche angestellt, betreut die Kinder im Alltag |
    | Joris Harms | 28.09.1980 | Antragsgegner | Geschäftsführer und Alleingesellschafter der Harms Lichtbau GmbH |
    | Tilda Harms | 11.11.2011 | Kind | 9. Klasse Gymnasium, Nachhilfe Mathematik |
    | Bent Harms | 04.05.2018 | Kind | 2. Klasse Grundschule, Hort, Fußball |

    Die Betreuung liegt seit Januar 2026 überwiegend bei Lene. Joris nimmt die Kinder nach der informellen Absprache jedes zweite Wochenende von Freitag 17:00 Uhr bis Sonntag 18:00 Uhr. Mittwochs holt er Bent zum Fußball und bringt ihn gegen 19:30 Uhr zurück. Tilda verweigert seit April einzelne Wochenenden, weil sie nach ihrer Aussage "nicht bei Papa im Büro schlafen" möchte.

    Mehrbedarf und Sonderbedarf sind bisher nicht sortiert. Die Nachhilfe von Tilda kostet 185,00 EUR monatlich. Bents Hort kostet nach Abzug des Zuschusses 158,00 EUR monatlich. Im Mai fiel eine Klassenfahrt-Anzahlung von 240,00 EUR an. Lene hat diese aus dem Dispo gezahlt.
    """)

    write_text(base / "03_zahlungsfluss_und_verzug.md", """
    # Zahlungsfluss und Verzug

    Lene forderte Joris am 12.02.2026 per E-Mail auf, Auskunft über sein Einkommen aus selbstständiger Tätigkeit, Geschäftsführergehalt, Privatentnahmen, Firmenwagen und Vermietungseinnahmen zu erteilen. Die E-Mail enthielt noch keine vollständig bezifferte Unterhaltsforderung, aber die Formulierung, dass Kindes- und Trennungsunterhalt "ab Februar 2026 in gesetzlicher Höhe" verlangt werde.

    Am 02.03.2026 zahlte Joris 1.500,00 EUR und schrieb im Betreff "mehr geht nicht". Ab April überwies er nur noch 900,00 EUR. Er legte eine BWA Januar bis März 2026 vor, aber keine Jahresabschlüsse 2023 bis 2025, keine Steuerbescheide, keine Summen- und Saldenlisten, keinen Anstellungsvertrag und keine Darlehensverträge.

    Für den Stufenantrag soll geprüft werden, ob ab Februar 2026 Verzug nach § 1613 BGB für Kindesunterhalt und über § 1361 Abs. 4 Satz 4 BGB i.V.m. § 1613 BGB für Trennungsunterhalt begründet ist. Für die einstweilige Anordnung kommt eine vorläufige Bezifferung auf Grundlage des sichtbaren Lebensstandards in Betracht; die Belege sind aber dünn.
    """)

    write_text(base / "04_einkommen_lenes_arbeitgeberbescheinigung.md", """
    # Einkommen Lene

    Lene legte Gehaltsabrechnungen Januar bis Mai 2026 vor. Das durchschnittliche Netto nach Abzug der gesetzlichen Abgaben beträgt 2.386,40 EUR. In März und April waren jeweils 80,00 EUR steuerfreie Fahrtkostenzuschüsse enthalten. Sie zahlt monatlich 46,00 EUR berufsbedingte ÖPNV-Kosten, außerdem 38,00 EUR Gewerkschaftsbeitrag. Eine private Altersvorsorge mit 110,00 EUR monatlichem Beitrag besteht seit 2017.

    Die Arbeitgeberbescheinigung weist keine Boni und keine Überstundenvergütung aus. Lene hat wegen der Betreuung ihre Arbeitszeit 2023 von 35 auf 30 Stunden reduziert. Eine kurzfristige Ausweitung auf Vollzeit ist nach Auskunft der Personalabteilung "erst nach dem Museumsumzug 2027 realistisch".
    """)

    write_text(base / "05_joris_unvollstaendige_auskunft.md", """
    # Unvollständige Auskunft Joris

    Joris übersandte am 22.03.2026 eine E-Mail mit dem Satz: "Meine BWA sagt alles, ich bin praktisch pleite." Angehängt war eine BWA Januar bis März 2026 der Harms Lichtbau GmbH. Das Geschäftsführergehalt wird dort mit 3.000,00 EUR monatlich verbucht. Gleichzeitig zeigt die BWA Kfz-Kosten von 1.284,00 EUR monatlich, Reisekosten von 2.918,00 EUR im März und "sonstige betriebliche Aufwendungen" von 18.440,00 EUR im Quartal.

    Nicht vorgelegt wurden: Steuerbescheide 2023 bis 2025, Einkommensteuererklärungen, Jahresabschlüsse der GmbH, Summen- und Saldenlisten, Gesellschafterdarlehen, Kontoauszüge Privatkonto, Nachweise zu Privatnutzung Pkw, Angaben zur Vermietung der Ferienwohnung Grömitz und der Geschäftsführer-Anstellungsvertrag. Lene berichtet, Joris habe im April in einem Chat geschrieben, er werde "nicht noch seine komplette Firma offenlegen".

    Das ist der Kern für die Stufenklage: Ohne vollständige Auskunft lässt sich der Unterhalt nicht seriös berechnen. Die einstweilige Anordnung soll nur eine vorläufige Sicherung leisten.
    """)

    write_text(base / "06_bwa_fragmente_firmenwagen_entnahmen.md", """
    # BWA-Fragmente und Privatnutzung

    Die BWA der Harms Lichtbau GmbH weist für Januar bis März 2026 einen vorläufigen Verlust von 14.870,00 EUR aus. Gleichzeitig sind die Umsatzerlöse gegenüber Q1 2025 nur um 8,2 % gefallen. Der größte Unterschied liegt in den Kostenpositionen Kfz, Fremdleistungen, Reisekosten und sonstige betriebliche Aufwendungen.

    Auffällig ist der Tesla Model Y, den Joris seit Oktober 2024 privat nutzt. Lene hat Fotos aus dem Familienurlaub 2025, auf denen der Wagen vor der Ferienwohnung steht. In den Gehaltsabrechnungen Joris ist ein geldwerter Vorteil nicht erkennbar; es liegt aber nur ein Screenshot einer Abrechnung vor.

    Außerdem findet sich in der BWA eine Buchung "Darlehen Gesellschafter JH Rückführung 9.500,00 EUR". Ohne Darlehensvertrag und Kontobeleg ist unklar, ob es sich um echte Rückzahlung, verdeckte Entnahme oder Liquiditätsverschiebung handelt.
    """)

    write_text(base / "07_ferienwohnung_groemitz_hinweise.md", """
    # Ferienwohnung Grömitz

    Lene fand über eine Ferienplattform ein Inserat "Lichtkoje Grömitz - Hafenblick und Sauna". Die Fotos zeigen nach ihrer Erinnerung Möbel aus der Wohnung von Joris' Mutter. Die Gastgeberantwort auf eine Anfrage kam vom Account "J. Harms". Der Kalender weist für Juli und August 2026 nur noch drei freie Wochen aus. Preise laut Screenshot: 148,00 EUR bis 219,00 EUR pro Nacht zuzüglich Reinigung.

    Es liegt kein Grundbuchauszug, kein Mietvertrag und keine Abrechnung vor. Joris behauptet, die Wohnung gehöre seiner Mutter und er helfe nur technisch. Lene hält das für falsch, weil Joris schon 2024 von "meiner kleinen Ostsee-Rente" gesprochen habe. Für die Auskunftsstufe sollten Eigentum, Nutzungsrecht, Einnahmen, Plattformkonto und Kontozufluss abgefragt werden.
    """)

    write_text(base / "08_auskunftsaufforderung_entwurf.md", """
    # Entwurf Auskunftsaufforderung

    Sehr geehrter Herr Harms,

    namens und in Vollmacht von Frau Lene Harms fordere ich Sie auf, zur Berechnung des Kindesunterhalts für Tilda und Bent sowie des Trennungsunterhalts für Frau Harms vollständige Auskunft über Ihre Einkünfte und Ihr Vermögen zu erteilen und die Angaben zu belegen.

    Für die selbstständige und gesellschaftsrechtlich vermittelte Tätigkeit sind insbesondere vorzulegen: Einkommensteuerbescheide und -erklärungen 2023 bis 2025, Jahresabschlüsse und BWA/SuSa der Harms Lichtbau GmbH 2023 bis laufend, Geschäftsführer-Anstellungsvertrag, Gehaltsabrechnungen, Gesellschafterdarlehensverträge, Kontoauszüge zu Entnahmen und Rückführungen, Nachweise zur privaten Nutzung des Tesla Model Y, Aufstellungen über Reisekosten und sonstige betriebliche Aufwendungen sowie sämtliche Unterlagen zur Ferienwohnung Grömitz einschließlich Plattformabrechnungen.

    Die Auskunft wird bis zum 28.06.2026 erwartet. Nach fruchtlosem Fristablauf wird gerichtliche Hilfe in Anspruch genommen, einschließlich Stufenantrag und Antrag auf vorläufigen Unterhalt.
    """)

    write_text(base / "09_stufenantrag_entwurf.md", """
    # Entwurf Stufenantrag

    An das Amtsgericht Hamburg - Familiengericht -

    In der Familiensache Lene Harms gegen Joris Harms wegen Kindesunterhalt und Trennungsunterhalt beantragen wir im Wege des Stufenantrags:

    1. Der Antragsgegner wird verpflichtet, der Antragstellerin Auskunft über seine Einkünfte aus nichtselbstständiger Arbeit, selbstständiger Tätigkeit, Geschäftsführervergütung, Gewinnausschüttungen, Privatentnahmen, geldwerten Vorteilen und Vermietung für die Jahre 2023, 2024, 2025 und das laufende Jahr 2026 zu erteilen.

    2. Der Antragsgegner wird verpflichtet, die Auskunft durch Vorlage der im Antrag einzeln bezeichneten Belege zu belegen.

    3. Der Antragsgegner wird verpflichtet, die Richtigkeit und Vollständigkeit der Auskunft an Eides statt zu versichern, falls die erteilte Auskunft unvollständig oder nicht mit hinreichender Sorgfalt erstellt erscheint.

    4. Nach Erteilung der Auskunft bleibt die Bezifferung von Kindesunterhalt, Trennungsunterhalt, Rückständen und laufendem Unterhalt vorbehalten.

    Begründung: Die Beteiligten sind verheiratet und leben seit dem 18.01.2026 getrennt. Die gemeinsamen minderjährigen Kinder leben im Haushalt der Antragstellerin. Der Antragsgegner hat seine Zahlungen ohne nachvollziehbare Einkommensdarlegung reduziert. Die bisher vorgelegte BWA reicht zur Unterhaltsberechnung nicht aus, weil sie nur ein Quartal betrifft und wesentliche Einkommensbestandteile aus Gesellschaft, Privatnutzung und Vermietung nicht offenlegt.
    """)

    write_text(base / "10_einstweilige_anordnung_vorlaeufiger_unterhalt.md", """
    # Einstweilige Anordnung: vorläufiger Unterhalt

    Für die Zwischenzeit bis zur Auskunftserteilung kommt ein Antrag auf einstweilige Anordnung in Betracht. Ziel ist nicht die endgültige Unterhaltsberechnung, sondern die Sicherung der laufenden Miete, des Kindesbedarfs und der Kranken-/Schulkosten.

    Vorläufige Bezifferung aus Mandantensicht:

    | Position | Betrag |
    |---|---:|
    | bisherige Gesamtzahlung Joris Februar | 2.250,00 EUR |
    | aktuelle Zahlung seit April | 900,00 EUR |
    | Warmmiete Ehewohnung | 1.780,00 EUR |
    | Hort Bent | 158,00 EUR |
    | Nachhilfe Tilda | 185,00 EUR |
    | geschätzte Unterdeckung monatlich | 1.100,00 EUR bis 1.450,00 EUR |

    Die Akte enthält keine sichere Endberechnung. Das System soll deshalb einen vorsichtigen EA-Antrag formulieren, der Mindestunterhalt der Kinder, Mehrbedarf und einen vorläufigen Trennungsunterhalt verlangt, ohne die spätere Stufenklage zu präjudizieren.
    """)

    write_text(base / "11_belegmatrix_luecken_und_beweiswert.md", """
    # Belegmatrix, Lücken und Beweiswert

    | Thema | Vorhanden | Fehlt | Bedeutung |
    |---|---|---|---|
    | Einkommen Lene | Gehaltsabrechnungen Jan-Mai 2026, Arbeitgeberbescheinigung | Steuerbescheid 2025 | solide Basis für Bedarf/Quote |
    | Einkommen Joris | BWA Q1 2026, Screenshot Gehalt März | Steuerbescheide, Jahresabschlüsse, Anstellungsvertrag | Auskunftsstufe notwendig |
    | Firmenwagen | Fotos, BWA-Kfz-Kosten | Nutzungsvertrag, geldwerter Vorteil | mögliches Einkommen |
    | Ferienwohnung | Plattformhinweis, Chat | Eigentum, Abrechnung, Konto | mögliches Vermietungseinkommen |
    | Zahlungen | Kontoauszug Jan-Jun 2026 | vollständiges Konto Lene | Rückstand und Verzug |
    | Kinderkosten | Hortvertrag, Nachhilfevertrag, Klassenfahrt | Zahlungsbelege Mai/Juni | Mehrbedarf/Sonderbedarf |
    """)

    write_text(base / "12_whatsapp_elternchat_unterhalt.md", """
    # Chat-Auszug Elternkommunikation

    **Joris, 07.04.2026, 22:18 Uhr:** Ich zahle jetzt 900. Mehr geht nicht. Firma ist im Minus.

    **Lene, 07.04.2026, 22:26 Uhr:** Dann schick mir bitte die Unterlagen, nicht nur die eine BWA. Ich kann die Miete nicht aus Luft bezahlen.

    **Joris, 07.04.2026, 22:31 Uhr:** Du bekommst keine komplette Firmenakte. Das ist mein Laden, nicht deiner.

    **Lene, 08.04.2026, 07:04 Uhr:** Es geht um Tilda und Bent. Hort, Nachhilfe, Klassenfahrt, Essen. Du warst letzte Woche mit dem Tesla in Grömitz und erzählst was von pleite.

    **Joris, 08.04.2026, 07:18 Uhr:** Grömitz geht dich nichts an. Das ist Familie.
    """)

    write_eml(
        base / "emails/2026-02-16_lenes_erstanfrage.eml",
        "Trennung - ich brauche Hilfe wegen Unterhalt",
        "Lene Harms <lene.harms@example.invalid>",
        "kanzlei@example.invalid",
        "Mon, 16 Feb 2026 08:42:00 +0100",
        """
        Guten Morgen,

        mein Mann ist im Januar ausgezogen. Er hat im Februar noch Geld überwiesen, sagt jetzt aber, dass seine Firma schlecht läuft. Ich glaube das nicht, weil er weiter Firmenwagen fährt und von Grömitz Einnahmen hat. Ich brauche bitte Hilfe, damit die Kinder und ich nicht in Rückstand geraten.

        Ich kann Gehaltsabrechnungen, Kontoauszüge und seine E-Mails mitbringen.

        Freundliche Grüße
        Lene Harms
        """,
    )
    write_eml(
        base / "emails/2026-03-08_joris_kuerzt_zahlung.eml",
        "Unterhalt März",
        "Joris Harms <joris.harms@example.invalid>",
        "Lene Harms <lene.harms@example.invalid>",
        "Sun, 8 Mar 2026 21:11:00 +0100",
        """
        Lene,

        ich überweise diesen Monat 1.500 EUR. Ab April werden es 900 EUR sein. Meine BWA ist negativ, ich kann nicht zaubern. Bitte komm mir nicht mit Anwälten wegen jeder Ausgabe, die Firma muss überleben.

        Joris
        """,
    )

    write_docx(
        base / "docx/auskunftsaufforderung_joris_harms.docx",
        "Aufforderung zur Auskunft und Belegvorlage",
        [
            ("Sachverhalt", ["Die Beteiligten leben seit dem 18.01.2026 getrennt. Die gemeinsamen Kinder Tilda und Bent leben im Haushalt von Frau Harms. Herr Harms hat seine laufenden Zahlungen seit April 2026 auf 900,00 EUR reduziert."]),
            ("Anforderung", ["Herr Harms wird aufgefordert, bis zum 28.06.2026 vollständige Auskunft über seine Einkünfte aus Geschäftsführertätigkeit, Gewinnausschüttungen, Privatentnahmen, Firmenwagennutzung und Vermietung zu erteilen und die bezeichneten Belege vorzulegen."]),
            ("Folgen", ["Nach fruchtlosem Fristablauf wird gerichtliche Auskunft im Wege des Stufenantrags sowie vorläufiger Unterhalt geltend gemacht."]),
        ],
    )
    write_docx(
        base / "docx/stufenantrag_unterhalt_harms_entwurf.docx",
        "Stufenantrag Unterhalt Harms - Entwurf",
        [
            ("Anträge", ["Der Antragsgegner wird zur Auskunft, Belegvorlage und gegebenenfalls Versicherung an Eides statt verpflichtet. Die Bezifferung des Kindes- und Trennungsunterhalts bleibt nach Auskunftserteilung vorbehalten."]),
            ("Begründung", ["Die bisherige Teil-BWA ist zur Ermittlung des unterhaltsrechtlich relevanten Einkommens eines selbstständigen Geschäftsführers unzureichend. Die Zahlungsreduzierung ist nicht nachvollziehbar."]),
            ("Eilbedürftigkeit", ["Die Antragstellerin trägt Miete, Hort, Nachhilfe und laufenden Kindesbedarf überwiegend allein. Die bisherigen Zahlungen reichen zur Deckung des notwendigen Bedarfs nicht aus."]),
        ],
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Schaetzung"
    ws.append(["Position", "Monat EUR", "Status", "Kommentar"])
    rows = [
        ["Netto Lene", 2386.40, "belegt", "Durchschnitt Jan-Mai 2026"],
        ["Joris GF-Gehalt", 3000.00, "teilbelegt", "nur Screenshot/BWA"],
        ["Privatnutzung Tesla", 650.00, "geschätzt", "geldwerter Vorteil offen"],
        ["Vermietung Grömitz", 900.00, "streitig", "Plattformindiz"],
        ["Sonstige Entnahmen", 1200.00, "streitig", "BWA-Position unklar"],
        ["Zahlung aktuell", 900.00, "belegt", "seit April"],
        ["Hort Bent", 158.00, "belegt", "Mehrbedarf"],
        ["Nachhilfe Tilda", 185.00, "belegt", "Mehrbedarf"],
    ]
    for row in rows:
        ws.append(row)
    workbook_style(ws)
    ws2 = wb.create_sheet("Luecken")
    ws2.append(["Unterlage", "Zeitraum", "Warum wichtig", "Status"])
    for row in [
        ["ESt-Bescheid Joris", "2023-2025", "Einkommen Drei-Jahres-Schnitt", "fehlt"],
        ["Jahresabschluss GmbH", "2023-2025", "Gewinn/Entnahmen", "fehlt"],
        ["SuSa/BWA", "2025-2026", "laufende Entwicklung", "Q1 2026 teilweise"],
        ["Firmenwagennutzung", "2024-2026", "geldwerter Vorteil", "fehlt"],
        ["Ferienwohnung", "2024-2026", "Vermietungseinkommen", "streitig"],
    ]:
        ws2.append(row)
    workbook_style(ws2)
    wb.save(base / "xlsx/vorlaeufige_unterhaltsschaetzung_harms.xlsx")

    write_csv(base / "csv/kontoauszug_unterhalt_zahlungen.csv", [
        ["Datum", "Buchungstext", "Betrag EUR", "Bemerkung"],
        ["2026-02-03", "Joris Unterhalt Februar", "2250.00", "vor Kürzung"],
        ["2026-03-02", "Joris Unterhalt März", "1500.00", "Ankündigung Engpass"],
        ["2026-04-02", "Joris Unterhalt April", "900.00", "gekürzt"],
        ["2026-05-04", "Joris Unterhalt Mai", "900.00", "gekürzt"],
        ["2026-06-03", "Joris Unterhalt Juni", "900.00", "gekürzt"],
    ])

    write_pdf(base / "pdfs/bwa_auszug_harms_lichtbau_q1_2026.pdf", "BWA-Auszug Harms Lichtbau GmbH Q1 2026", [
        ("Auszug", [["Position", "Januar", "Februar", "März"], ["Umsatz", "58.400", "49.900", "61.700"], ["GF-Gehalt", "3.000", "3.000", "3.000"], ["Kfz-Kosten", "1.184", "1.209", "1.459"], ["Reisekosten", "612", "774", "2.918"], ["Sonstige Aufwendungen", "5.980", "4.112", "8.348"]]),
        ("Notiz", "Der Auszug wurde von Joris Harms ohne Anlagen übersandt. Kontennachweise, Summen- und Saldenliste und Vorjahresvergleich fehlen."),
    ])
    write_pdf(base / "pdfs/arbeitgeberbescheinigung_lene_harms.pdf", "Arbeitgeberbescheinigung Lene Harms", [
        ("Beschäftigung", "Frau Lene Harms ist seit 01.09.2017 bei der Nordkai Museum Service gGmbH beschäftigt. Seit 01.08.2023 beträgt die regelmäßige Wochenarbeitszeit 30 Stunden."),
        ("Entgelt", [["Monat", "Netto EUR"], ["Januar 2026", "2.382,10"], ["Februar 2026", "2.361,84"], ["März 2026", "2.405,22"], ["April 2026", "2.398,17"], ["Mai 2026", "2.384,67"]]),
    ])

    write_text(base / "rubric.yaml", """
    expected_use:
      - Auskunftsanspruch und Belegvorlage erkennen
      - Stufenantrag nach Auskunft, Versicherung und Zahlung strukturieren
      - Vorläufigen Unterhalt nur vorsichtig schätzen
      - Lücken bei Selbstständigen-Einkommen und Privatnutzung benennen
    """)


def build_calculation_case() -> None:
    slug = "unterhalt-berechnungsakte-vollstaendig-rosenheim"
    base = TESTAKTEN / slug
    for d in ["emails", "docx", "xlsx", "pdfs", "csv", "gesamt-pdf"]:
        (base / d).mkdir(parents=True, exist_ok=True)

    write_text(base / "README.md", """
    # Akte: Unterhalt Mira Aiblinger ./. Tobias Aiblinger — vollständige Zahlen, Trennungsunterhalt, nacheheliche Varianten

    **Arbeitsakte.** Alle Personen, Arbeitgeber, Anschriften und Aktenzeichen sind anonymisiert. Die Akte gehört fachlich zum Plugin `fachanwalt-familienrecht`.

    ## Kurzbild

    Mira und Tobias Aiblinger leben seit dem 01.02.2026 getrennt. Die Ehe wurde 2014 geschlossen. Die Kinder Klara (16), Emil (12) und Nino (6) leben überwiegend bei Mira in Rosenheim; Tobias betreut jedes zweite Wochenende und einen Nachmittag pro Woche. Alle Gehaltsabrechnungen, Steuerbescheide, Versicherungen, Kreditraten, Wohnkosten und Kinderkosten liegen vor. Die Akte soll das System zu einer vollständigen Unterhaltsberechnung bringen: Kindesunterhalt 2026, Mehrbedarf/Sonderbedarf, Trennungsunterhalt, nachehelicher Aufstockungs-/Betreuungsunterhalt, Varianten mit Steuerklasse, Wohnvorteil und Erwerbsausweitung.

    Anders als die Stufenklage-Akte ist diese Akte zahlenreich und grundsätzlich berechenbar. Streit besteht nicht über das Ob der Auskunft, sondern über den Rechenweg und die Varianten.

    ## Aktenstücke

    | Nr. | Datei | Inhalt |
    |---|---|---|
    | 01 | `01_sachverhalt_und_trennung.md` | Chronologie, Rollen, Trennungsdatum |
    | 02 | `02_kinder_betreuung_mehrbedarf.md` | Kinder, Betreuung, Mehrbedarf, Sonderbedarf |
    | 03 | `03_einkommen_tobias.md` | Einkommen, Bonus, Dienstwagen, Altersvorsorge |
    | 04 | `04_einkommen_mira.md` | Teilzeit, mögliche Erwerbsausweitung, Krankenversicherung |
    | 05 | `05_bereinigungen_schulden_versicherungen.md` | berufsbedingte Aufwendungen, Kredit, Altersvorsorge |
    | 06 | `06_wohnwert_und_hauslasten.md` | Haus, Wohnvorteil, Zins/Tilgung, Nebenkosten |
    | 07 | `07_kindesunterhalt_2026_datenblatt.md` | Düsseldorfer Tabelle 2026, Kindergeld 259 EUR |
    | 08 | `08_trennungsunterhalt_rechenvarianten.md` | Varianten mit Erwerbstätigenbonus |
    | 09 | `09_nachehelicher_unterhalt_varianten.md` | Aufstockung, Betreuung, Befristung, Herabsetzung |
    | 10 | `10_vergleichsangebot_entwurf.md` | Vergleichsvorschlag mit Zahlbeträgen und Anpassungsklausel |
    | 11 | `11_rechenhinweise_kanzlei.md` | Bearbeiterhinweise ohne fertige Lösung |
    | 12 | `12_emailverkehr_zahlenfreigabe.md` | E-Mail-Auszüge zur Zahlenabstimmung |

    ## Anhänge

    | Datei | Inhalt |
    |---|---|
    | `xlsx/unterhaltsrechner_aiblinger_2026.xlsx` | Berechnungsmodell mit Eingaben, Kindesunterhalt, Ehegattenunterhalt, Varianten |
    | `docx/vergleichsentwurf_unterhalt_aiblinger.docx` | Vergleichsentwurf |
    | `docx/mandantenbrief_berechnungserlaeuterung.docx` | Erläuterung an Mandantin |
    | `pdfs/gehaltsuebersicht_tobias_2025_2026.pdf` | Lohn- und Bonusübersicht Tobias |
    | `pdfs/gehaltsuebersicht_mira_2025_2026.pdf` | Lohnübersicht Mira |
    | `csv/kinderkosten_2026.csv` | Mehrbedarf/Sonderbedarf |
    | `emails/2026-05-18_zahlenfreigabe_tobias_anwalt.eml` | Freigabe Einkommenszahlen |
    | `emails/2026-05-21_mira_fragt_varianten.eml` | Mandantin fragt Varianten |
    """)

    write_text(base / "01_sachverhalt_und_trennung.md", """
    # Sachverhalt und Trennung

    Mira Aiblinger und Tobias Aiblinger heirateten am 19.07.2014. Die Trennung erfolgte nach übereinstimmender Darstellung am 01.02.2026. Tobias zog in eine Mietwohnung in Kolbermoor. Mira blieb mit den Kindern im Haus in Rosenheim. Ein Scheidungsantrag soll frühestens im Februar 2027 gestellt werden.

    Tobias ist Entwicklungsleiter bei der Berglicht Robotics GmbH. Mira ist angestellte Physiotherapeutin mit 24 Wochenstunden. Sie möchte ab September 2026 auf 30 Wochenstunden erhöhen, falls die Betreuung von Nino gesichert ist. Tobias meint, Mira könne sofort voll arbeiten. Mira verweist auf Ninos Einschulung, Emil mit ADHS-Diagnostik und die Nachmittagslogistik.
    """)

    write_text(base / "02_kinder_betreuung_mehrbedarf.md", """
    # Kinder, Betreuung und Mehrbedarf

    | Kind | Alter 2026 | Lebensmittelpunkt | Besonderheiten |
    |---|---:|---|---|
    | Klara | 16 | Mira | Gymnasium, Musikschule, Zahnspange abgeschlossen |
    | Emil | 12 | Mira | ADHS-Diagnostik, Lerntherapie |
    | Nino | 6 | Mira | Einschulung September 2026, Hortplatz zugesagt |

    Tobias betreut jedes zweite Wochenende und mittwochs von 15:30 Uhr bis 19:00 Uhr. Eine paritätische Betreuung liegt nicht vor. Der Mehrbedarf ist zwischen den Eltern quotal zu verteilen, soweit er angemessen und belegt ist. Klara hat Musikschulkosten von 86,00 EUR monatlich. Emil erhält Lerntherapie für 210,00 EUR monatlich. Nino hat ab September Hortkosten von voraussichtlich 145,00 EUR monatlich.
    """)

    write_text(base / "03_einkommen_tobias.md", """
    # Einkommen Tobias

    Tobias erzielte 2025 ein durchschnittliches monatliches Netto von 6.840,00 EUR einschließlich anteiligem Bonus. Für 2026 liegen die Abrechnungen Januar bis Mai vor. Das Monatsnetto ohne Bonus beträgt 6.210,00 EUR. Im April 2026 wurde ein Bonus für 2025 in Höhe von netto 9.600,00 EUR gezahlt. Unterhaltsrechtlich ist eine Zwölftelung zu prüfen.

    Zusätzlich nutzt Tobias einen Dienstwagen Audi Q5. Der geldwerte Vorteil ist in der Abrechnung mit 685,00 EUR brutto ausgewiesen; der tatsächliche unterhaltsrechtliche Vorteil ist streitig, weil Tobias nach dem Auszug täglich weiter nach München pendelt. Er zahlt 320,00 EUR monatlich in eine zusätzliche Altersvorsorge und 72,00 EUR Berufsunfähigkeitsversicherung.
    """)

    write_text(base / "04_einkommen_mira.md", """
    # Einkommen Mira

    Mira erzielt bei 24 Wochenstunden ein durchschnittliches Monatsnetto von 2.180,00 EUR. Die Praxisleitung hat eine Erhöhung auf 30 Wochenstunden ab September 2026 angeboten; erwartetes Netto dann 2.610,00 EUR. Eine sofortige Vollzeitstelle ist nicht angeboten und wegen Ninos Einschulung sowie Emils Therapie organisatorisch nicht gesichert.

    Mira zahlt 95,00 EUR monatlich private Zusatzversicherung für die Kinder, 64,00 EUR berufsbedingte Fahrtkosten und 130,00 EUR Altersvorsorge. Der Arbeitgeberzuschuss zur Fortbildung ist einmalig und nicht als dauerhaftes Einkommen geplant.
    """)

    write_text(base / "05_bereinigungen_schulden_versicherungen.md", """
    # Bereinigungen, Schulden und Versicherungen

    Tobias zahlt für die Ehewohnung keinen Kredit mehr, weil Mira die Hauslasten seit Februar trägt. Er zahlt Miete von 1.420,00 EUR warm. Seine private Altersvorsorge von 320,00 EUR monatlich besteht seit 2016. Die Berufsunfähigkeitsversicherung ist berufsbezogen plausibel. Streitig ist ein Konsumentenkredit über 310,00 EUR monatlich für ein Motorrad, das Tobias nach der Trennung gekauft hat.

    Mira trägt Zins und Tilgung für das Hausdarlehen mit 1.480,00 EUR monatlich. Davon entfallen 780,00 EUR auf Zinsen und 700,00 EUR auf Tilgung. Die Tilgung ist beim Wohnvorteil gesondert zu bewerten. Nebenkosten des Hauses betragen 510,00 EUR monatlich. Der objektive Wohnwert wird von beiden Seiten vorläufig mit 1.850,00 EUR angesetzt; Mira meint, wegen Trennung und nicht abgeschlossener Vermögensauseinandersetzung sei zunächst nur ein angemessener Wohnwert anzusetzen.
    """)

    write_text(base / "06_wohnwert_und_hauslasten.md", """
    # Wohnwert und Hauslasten

    Das Haus steht im hälftigen Miteigentum. Mira nutzt es mit den drei Kindern. Tobias verlangt, den vollen objektiven Wohnwert anzusetzen. Mira möchte bis zur Scheidung nur den angemessenen Wohnwert ansetzen, weil ein Verkauf oder Auszug nicht kurzfristig zumutbar sei. Die Akte enthält deshalb zwei Varianten:

    1. Variante A: objektiver Wohnwert 1.850,00 EUR, Zinslast 780,00 EUR, keine volle Tilgungsberücksichtigung.
    2. Variante B: angemessener Wohnwert 1.250,00 EUR bis August 2026, ab September 2026 stufenweise Erhöhung.

    Für das System ist wichtig, Wohnvorteil, Kindesunterhalt und Ehegattenunterhalt nicht doppelt zu verrechnen.
    """)

    write_text(base / "07_kindesunterhalt_2026_datenblatt.md", """
    # Kindesunterhalt 2026: Datenblatt

    Die Akte arbeitet mit der Düsseldorfer Tabelle Stand 01.01.2026. Kindergeld beträgt 259,00 EUR je Kind. Die konkreten Tabellenbeträge sollen vom System anhand aktueller Quelle überprüft werden; in der Excel-Datei sind Arbeitswerte eingetragen, nicht als amtliche Tabelle gemeint.

    | Kind | Altersstufe | Lebensmittelpunkt | Kindergeldbezug | Mehrbedarf |
    |---|---|---|---|---:|
    | Klara | 12-17 | Mira | Mira | 86,00 EUR Musikschule |
    | Emil | 12-17 | Mira | Mira | 210,00 EUR Lerntherapie |
    | Nino | 0-5 bis August, 6-11 ab September | Mira | Mira | 145,00 EUR Hort ab September |

    Die Betreuung Tobias' bleibt deutlich unter Wechselmodell. Ein erweitertes Umgangsmodell kann Auswirkungen auf Mehrbedarf und einzelne Fahrtkosten haben, aber nicht automatisch auf die Barunterhaltspflicht.
    """)

    write_text(base / "08_trennungsunterhalt_rechenvarianten.md", """
    # Trennungsunterhalt: Rechenvarianten

    Variante 1 rechnet mit Tobias' durchschnittlichem Jahresnetto einschließlich Bonus und Miras aktuellem Einkommen. Variante 2 rechnet ab September 2026 mit Miras 30-Stunden-Netto. Variante 3 setzt beim Wohnwert nur einen angemessenen Wohnwert an. Variante 4 setzt den objektiven Wohnwert vollständig an.

    Der Erwerbstätigenbonus, Kindesunterhalt als Vorwegabzug, berücksichtigungsfähige Altersvorsorge, Wohnwert und Schulden sind sichtbar getrennt. Der Output soll nicht nur eine Zahl liefern, sondern zeigen, welche Parameter den Zahlbetrag bewegen und welche Variante verhandlungsfähig ist.
    """)

    write_text(base / "09_nachehelicher_unterhalt_varianten.md", """
    # Nachehelicher Unterhalt: Varianten

    Für die Zeit nach Scheidung stehen drei Linien im Raum:

    1. Betreuungsunterhalt wegen Nino und Emils Förderbedarf. Mira argumentiert, eine sofortige Vollzeittätigkeit sei bis mindestens Sommer 2028 unzumutbar.
    2. Aufstockungsunterhalt wegen ehebedingter Teilzeit und Einkommensdifferenz. Tobias argumentiert, Mira habe ihren Beruf nie aufgegeben und könne aufstocken.
    3. Befristung und Herabsetzung. Tobias will eine klare Befristung nach zwei Jahren; Mira verlangt eine Anpassungsklausel nach tatsächlicher Erwerbsausweitung und Betreuungslage.

    Die Akte enthält genug Zahlen, um Korridore zu bilden, aber keine endgültige gerichtliche Bewertung.
    """)

    write_text(base / "10_vergleichsangebot_entwurf.md", """
    # Vergleichsangebot Entwurf

    Tobias zahlt ab 01.07.2026 Kindesunterhalt nach der jeweils einschlägigen Einkommensgruppe der Düsseldorfer Tabelle, dynamisiert in Höhe von 120 % des Mindestunterhalts, abzüglich hälftigen Kindergelds. Mehrbedarf für Lerntherapie, Musikschule und Hort wird nach Einkommensquote getragen.

    Tobias zahlt Trennungsunterhalt an Mira in monatlicher Höhe von 1.180,00 EUR bis einschließlich August 2026 und ab September 2026 in Höhe von 920,00 EUR, jeweils unter Vorbehalt der Neuberechnung nach Steuerklassenwechsel, Bonuszahlung 2026 und Wohnwertklärung.

    Für nachehelichen Unterhalt vereinbaren die Beteiligten zunächst eine Evaluationsklausel sechs Monate nach Rechtskraft der Scheidung. Eine Befristung bleibt vorbehalten.
    """)

    write_text(base / "11_rechenhinweise_kanzlei.md", """
    # Rechenhinweise der Kanzlei

    Diese Akte soll nicht automatisch nur eine einzige Zahl auswerfen. Gewünscht ist ein Berechnungsblatt mit mindestens vier Varianten: aktuell, ab September mit höherem Einkommen Mira, objektiver Wohnwert, angemessener Wohnwert. Kindesunterhalt ist vor dem Ehegattenunterhalt zu berücksichtigen. Mehrbedarf der Kinder ist gesondert zu quoteln.

    Bitte im Ergebnis immer ausweisen: zugrunde gelegtes Einkommen, Bereinigungen, Quote, Kindergeld, Zahlbetrag, offene Streitpunkte und Anpassungsmechanismus. Wenn ein Tabellenwert nicht live verifiziert ist, muss der Output das klar markieren.
    """)

    write_text(base / "12_emailverkehr_zahlenfreigabe.md", """
    # E-Mail-Auszüge zur Zahlenfreigabe

    **18.05.2026, RA Tobias:** Wir geben die Einkommenszahlen 2025/2026 frei, bestehen aber darauf, den Bonus auf drei Jahre zu glätten. Der Dienstwagen ist im Netto bereits abgebildet.

    **21.05.2026, Mira:** Ich brauche eine Berechnung, die ich verstehe. Tobias sagt, ich solle Vollzeit arbeiten. Ich will wissen, was passiert, wenn ich ab September 30 Stunden mache, aber ich kann nicht jetzt schon Vollzeit versprechen.

    **24.05.2026, Kanzlei intern:** Bitte Variante mit angemessenem Wohnwert und Variante mit objektivem Wohnwert nebeneinander. Emils Lerntherapie nicht vergessen; Tobias hat dem bislang mündlich zugestimmt.
    """)

    write_eml(
        base / "emails/2026-05-18_zahlenfreigabe_tobias_anwalt.eml",
        "Aiblinger - Einkommensunterlagen Tobias",
        "ra.tobias@example.invalid",
        "kanzlei@example.invalid",
        "Mon, 18 May 2026 14:06:00 +0200",
        """
        Sehr geehrte Kollegin,

        anbei übersenden wir die Gehaltsübersicht unseres Mandanten einschließlich Bonus 2025 und der Abrechnungen Januar bis Mai 2026. Wir gehen davon aus, dass der Bonus nicht monatsgenau, sondern geglättet zu berücksichtigen ist. Zum Dienstwagen ist anzumerken, dass die Versteuerung bereits in der Abrechnung enthalten ist.

        Mit freundlichen kollegialen Grüßen
        RA Dr. Leitner
        """,
    )
    write_eml(
        base / "emails/2026-05-21_mira_fragt_varianten.eml",
        "Bitte Varianten erklären",
        "Mira Aiblinger <mira.aiblinger@example.invalid>",
        "kanzlei@example.invalid",
        "Thu, 21 May 2026 09:33:00 +0200",
        """
        Guten Morgen,

        ich verstehe die vielen Zahlen nicht. Können Sie bitte einmal rechnen, was Tobias für die Kinder zahlen muss, was für mich bis zur Scheidung herauskommt und was wäre, wenn ich ab September sechs Stunden mehr arbeite? Ich will keine überzogene Forderung, aber ich möchte wissen, was realistisch ist.

        Viele Grüße
        Mira
        """,
    )

    write_docx(
        base / "docx/vergleichsentwurf_unterhalt_aiblinger.docx",
        "Unterhaltsvergleich Aiblinger - Entwurf",
        [
            ("Kindesunterhalt", ["Der Kindesunterhalt wird dynamisiert nach der Düsseldorfer Tabelle 2026 gezahlt. Das Kindergeld wird nach den gesetzlichen Regeln angerechnet. Mehrbedarf wird nach Einkommensquote getragen."]),
            ("Trennungsunterhalt", ["Tobias zahlt vorläufig Trennungsunterhalt in zwei Stufen: bis August 2026 nach aktuellem Einkommen Mira, ab September 2026 unter Berücksichtigung der geplanten Stundenaufstockung."]),
            ("Anpassung", ["Bonuszahlungen, Steuerklassenwechsel, Wohnwert und wesentliche Änderungen der Betreuung werden jeweils zum Folgemonat nach Nachweis neu berechnet."]),
        ],
    )
    write_docx(
        base / "docx/mandantenbrief_berechnungserlaeuterung.docx",
        "Mandantenbrief - Erläuterung der Unterhaltsvarianten",
        [
            ("Ausgangspunkt", ["Wir rechnen zunächst mit den belegten Nettoeinkommen, bereinigen berufsbedingte Aufwendungen und berücksichtigen den Kindesunterhalt vor dem Ehegattenunterhalt."]),
            ("Warum mehrere Varianten?", ["Der Wohnwert und Ihre mögliche Stundenaufstockung verändern den Zahlbetrag deutlich. Deshalb ist eine einzige Zahl für die Verhandlung weniger hilfreich als ein Korridor mit sauber benannten Annahmen."]),
            ("Nächster Schritt", ["Wir schlagen vor, den Vergleich mit einer Anpassungsklausel zu versehen, damit Bonus, Steuerklasse und Betreuungsentwicklung nicht jedes Mal ein neues Verfahren auslösen."]),
        ],
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Eingaben"
    ws.append(["Parameter", "Wert", "Kommentar"])
    for row in [
        ["Netto Tobias laufend", 6210.00, "Jan-Mai 2026 Durchschnitt ohne Bonus"],
        ["Bonus Tobias monatlich", 800.00, "9600 EUR / 12"],
        ["Netto Mira aktuell", 2180.00, "24 Std."],
        ["Netto Mira ab Sep", 2610.00, "30 Std. Angebot"],
        ["Wohnwert angemessen", 1250.00, "Variante Trennung"],
        ["Wohnwert objektiv", 1850.00, "Variante Streit"],
        ["Hauszins", 780.00, "monatlich"],
        ["Kindergeld je Kind", 259.00, "2026"],
    ]:
        ws.append(row)
    workbook_style(ws)
    ws2 = wb.create_sheet("Kindesunterhalt")
    ws2.append(["Kind", "Arbeitsbedarf EUR", "Kindergeldanteil", "Zahlbetrag EUR", "Mehrbedarf", "Quote offen"])
    for row in [
        ["Klara", 784.00, 129.50, 654.50, 86.00, "Quote Eltern"],
        ["Emil", 784.00, 129.50, 654.50, 210.00, "Quote Eltern"],
        ["Nino bis Aug", 584.00, 129.50, 454.50, 0.00, "Quote Eltern"],
        ["Nino ab Sep", 671.00, 129.50, 541.50, 145.00, "Quote Eltern"],
    ]:
        ws2.append(row)
    workbook_style(ws2)
    ws3 = wb.create_sheet("Varianten")
    ws3.append(["Variante", "Tobias bereinigt", "Mira bereinigt", "Wohnwert", "Kindesunterhalt vorweg", "Rechenkorridor Trennungsunterhalt", "Kommentar"])
    for row in [
        ["A aktuell angemessener Wohnwert", 7010.00, 2180.00 + 1250.00 - 780.00, 1250.00, 1763.50, 1180.00, "vor September"],
        ["B ab September angemessener Wohnwert", 7010.00, 2610.00 + 1250.00 - 780.00, 1250.00, 1850.50, 920.00, "30 Std."],
        ["C objektiver Wohnwert", 7010.00, 2180.00 + 1850.00 - 780.00, 1850.00, 1763.50, 940.00, "Streit Tobias"],
        ["D Bonus geglättet 3 Jahre", 6610.00, 2180.00 + 1250.00 - 780.00, 1250.00, 1763.50, 1030.00, "Verhandlung"],
    ]:
        ws3.append(row)
    workbook_style(ws3)
    wb.save(base / "xlsx/unterhaltsrechner_aiblinger_2026.xlsx")

    write_csv(base / "csv/kinderkosten_2026.csv", [
        ["Kind", "Kostenart", "Betrag EUR", "Frequenz", "Beleg"],
        ["Klara", "Musikschule", "86.00", "monatlich", "Vertrag 2025"],
        ["Emil", "Lerntherapie", "210.00", "monatlich", "Therapievereinbarung"],
        ["Nino", "Hort", "145.00", "monatlich ab Sep", "Platzbestätigung"],
        ["Klara", "Klassenfahrt", "390.00", "einmalig Okt", "Schulmail"],
    ])

    write_pdf(base / "pdfs/gehaltsuebersicht_tobias_2025_2026.pdf", "Gehaltsübersicht Tobias Aiblinger", [
        ("Laufendes Netto", [["Monat", "Netto EUR"], ["Jan 2026", "6.188"], ["Feb 2026", "6.205"], ["März 2026", "6.221"], ["April 2026", "15.811 inkl. Bonus"], ["Mai 2026", "6.236"]]),
        ("Bonus", "Im April 2026 wurde ein Bonus für das Geschäftsjahr 2025 in Höhe von netto 9.600 EUR gezahlt. Die rechtliche Behandlung ist zwischen den Beteiligten streitig."),
    ])
    write_pdf(base / "pdfs/gehaltsuebersicht_mira_2025_2026.pdf", "Gehaltsübersicht Mira Aiblinger", [
        ("Laufendes Netto", [["Monat", "Netto EUR"], ["Jan 2026", "2.168"], ["Feb 2026", "2.179"], ["März 2026", "2.184"], ["April 2026", "2.192"], ["Mai 2026", "2.177"]]),
        ("Arbeitszeit", "Aktuell 24 Wochenstunden. Angebot der Praxisleitung: 30 Wochenstunden ab September 2026 mit voraussichtlichem Netto von 2.610 EUR."),
    ])

    write_text(base / "rubric.yaml", """
    expected_use:
      - Kindesunterhalt und Mehrbedarf mit Tabellenstand 2026 prüfen
      - Trennungsunterhalt in mehreren Varianten berechnen
      - Nacheheliche Unterhaltslinien bilden
      - Wohnwert, Bonus und Erwerbsausweitung transparent behandeln
    """)


def main() -> None:
    build_stufenklage_case()
    build_calculation_case()
    print("Unterhalts-Testakten erzeugt.")


if __name__ == "__main__":
    main()
